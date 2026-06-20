"""
Módulo 7 — Geração do Relatório Word (ABNT NBR 14724) — Multi-Stack
Texto adaptado dinamicamente à tecnologia detectada, sem assumir WordPress.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from modules.logger import log

BLUE_HDR  = RGBColor(0x1F, 0x38, 0x64)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

SEV_ORDER = {"URGENTE": 0, "ALTO": 1, "MÉDIO": 2, "BAIXO": 3, "INFORMATIVO": 4, "DESCARTADA": 5}
def _sev_key(v): return SEV_ORDER.get(v.get("severity","").upper(), 99)


def _set_cell_bg(cell, color_hex):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _set_cell_borders(cell):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"4"); b.set(qn("w:space"),"0"); b.set(qn("w:color"),"BFBFBF")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _para(doc, text, size=12, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
           color=None, space_before=0, space_after=6, line_spacing=1.5, indent=True):
    p = doc.add_paragraph(); p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(space_before); fmt.space_after = Pt(space_after); fmt.line_spacing = line_spacing
    if indent: fmt.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    r.font.name = "Arial"; r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
    if color: r.font.color.rgb = color
    return p


def _heading(doc, text, level=1):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.space_before = Pt(18 if level==1 else 12); fmt.space_after = Pt(6); fmt.line_spacing = 1.5
    r = p.add_run(text.upper() if level==1 else text)
    r.font.name = "Arial"; r.font.size = Pt(12); r.font.bold = True
    return p


def _bullet(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5; p.paragraph_format.left_indent = Cm(1.25)
    r = p.add_run(f"– {text}")
    r.font.name = "Arial"; r.font.size = Pt(12)
    return p


def _page_break(doc): doc.add_page_break()


def _section_title(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12); p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = "Arial"; r.font.size = Pt(12); r.font.bold = True
    return p


def _add_table(doc, headers, rows, col_widths):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr_row = table.rows[0]
    for ci, (hdr, w) in enumerate(zip(headers, col_widths)):
        cell = hdr_row.cells[ci]; cell.width = Cm(w)
        _set_cell_bg(cell, "1F3864"); _set_cell_borders(cell)
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(hdr)
        r.font.name="Arial"; r.font.size=Pt(9); r.font.bold=True; r.font.color.rgb = WHITE
        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        shade = "D9E2F3" if ri % 2 == 0 else "FFFFFF"
        for ci, (val, w) in enumerate(zip(row_data, col_widths)):
            cell = row.cells[ci]; cell.width = Cm(w)
            _set_cell_bg(cell, shade if ci > 0 else "D9E2F3"); _set_cell_borders(cell)
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(str(val))
            r.font.name="Arial"; r.font.size=Pt(9); r.font.bold = (ci==0)
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    doc.add_paragraph()
    return table


def generate_report(results: dict, output_path):
    log("Inicializando documento Word...")
    doc = Document()

    for s in doc.sections:
        s.page_height=Cm(29.7); s.page_width=Cm(21.0)
        s.top_margin=Cm(3.0); s.bottom_margin=Cm(2.0); s.left_margin=Cm(3.0); s.right_margin=Cm(2.0)
    doc.styles["Normal"].font.name = "Arial"; doc.styles["Normal"].font.size = Pt(12)

    domain      = results.get("domain", "")
    team        = results.get("team", "[NOME DA EQUIPE]")
    members     = results.get("members", [])
    advisor     = results.get("advisor", "[Nome do Professor(a)]")
    institution = results.get("institution", "")
    course      = results.get("course", "")
    date        = results.get("date", datetime.now().strftime("%d/%m/%Y"))
    dns     = results.get("dns", {})
    hdrs    = results.get("headers", {})
    tech    = results.get("tech", {})
    ep      = results.get("endpoints", {})
    vs      = results.get("vulnscan", {})
    vuls    = sorted(results.get("vulnerabilities", []), key=_sev_key)

    tech_name = tech.get("primary_tech_name", "Tecnologia não identificada")
    tech_cat  = tech.get("primary_tech_category", "")
    tech_ver  = tech.get("tech_specific", {}).get("version", "")
    tech_full = f"{tech_name} {tech_ver}".strip() if tech_ver else tech_name

    altas  = sum(1 for v in vuls if v.get("severity","").upper() in ("ALTO","URGENTE"))
    medias = sum(1 for v in vuls if v.get("severity","").upper() == "MÉDIO")
    baixas = sum(1 for v in vuls if v.get("severity","").upper() == "BAIXO")
    infos  = sum(1 for v in vuls if v.get("severity","").upper() == "INFORMATIVO")
    desc   = sum(1 for v in vuls if v.get("severity","").upper() == "DESCARTADA")

    # ── CAPA ──────────────────────────────────────────────────────────────────
    log("Gerando capa...")
    if institution:
        _para(doc, institution, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=2)
    if course:
        _para(doc, course, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=60)
    elif institution:
        _para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=60)
    _para(doc, "RELATÓRIO TÉCNICO DE PENTEST", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=4)
    _para(doc, "Reconhecimento, Varredura e Classificação de Vulnerabilidades", italic=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=4)
    _para(doc, f"Alvo: {domain}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=60)
    nota_capa = (
        "Relatório apresentado como requisito acadêmico em Cybersecurity."
        if institution else
        "Relatório técnico de avaliação de segurança (pentest)."
    )
    _para(doc, nota_capa, italic=True, size=11, align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False, space_after=6)
    if team and team != "[NOME DA EQUIPE]":
        _para(doc, f"Equipe: {team}", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False, space_after=4)
    if members:
        for i, m in enumerate(members, 1):
            _para(doc, f"Integrante {i}: {m}", size=11, align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False, space_after=3)
    if advisor and advisor != "[Responsável Técnico]":
        _para(doc, f"Responsável: {advisor}", size=11, align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False, space_after=3)
    _para(doc, "", space_after=60, indent=False)
    report_year = date.split("/")[-1] if "/" in date else datetime.now().strftime("%Y")
    _para(doc, report_year, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    _page_break(doc)

    # ── RESUMO ────────────────────────────────────────────────────────────────
    log("Gerando resumo e abstract...")
    _section_title(doc, "RESUMO")
    resumo = (
        f"O presente relatório documenta o reconhecimento técnico e a varredura de vulnerabilidades realizados sobre o ambiente {domain}, "
        f"seguindo a metodologia OWASP Web Security Testing Guide (WSTG) v4.2. A identificação automática de tecnologias revelou que o "
        f"ambiente utiliza {tech_full} (categoria: {tech_cat or 'não classificada'}). Foram executados reconhecimento passivo (DNS, WHOIS, DNSSEC), "
        f"análise de cabeçalhos HTTP, fingerprinting de tecnologias, enumeração de endpoints específicos da stack identificada e varredura "
        f"automatizada de vulnerabilidades, complementados por verificações genéricas de segurança aplicáveis a qualquer tecnologia web."
    )
    _para(doc, resumo, indent=False, line_spacing=1.0)
    _para(doc, f"Foram identificadas {len(vuls)} vulnerabilidades preliminares, sendo {altas} de severidade Alta, {medias} Médias, {baixas} Baixas e {infos} Informativas, além de {desc} descartadas após verificação.", indent=False, line_spacing=1.0)
    _para(doc, "", indent=False, space_after=4)
    kw = doc.add_paragraph(); kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; kw.paragraph_format.line_spacing = 1.0
    r = kw.add_run(f"Palavras-chave: Cybersecurity. Pentest. OWASP WSTG. {tech_name}. Reconhecimento. Varredura.")
    r.font.name="Arial"; r.font.size=Pt(12); r.font.italic=True
    _page_break(doc)

    # ── SUMÁRIO ───────────────────────────────────────────────────────────────
    _section_title(doc, "SUMÁRIO")
    for line, pg in [
        ("1  INTRODUÇÃO","3"), ("2  IDENTIFICAÇÃO DO PROJETO","3"),
        ("3  METODOLOGIA","4"), ("4  RECONHECIMENTO E VARREDURA","4"),
        ("   4.1  Tecnologia Detectada","4"), ("   4.2  DNS e WHOIS","5"),
        ("   4.3  Cabeçalhos HTTP","5"), ("   4.4  Endpoints e Arquivos Sensíveis","6"),
        ("   4.5  Tabela Consolidada de Vulnerabilidades","7"),
        ("5  ANÁLISE DE IMPACTO","8"), ("6  SUGESTÕES DE CORREÇÃO","9"),
        ("7  CONSIDERAÇÕES FINAIS","10"), ("REFERÊNCIAS","11"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(3); p.paragraph_format.line_spacing=1.5
        r1 = p.add_run(line); r1.font.name="Arial"; r1.font.size=Pt(12)
        r2 = p.add_run(f"  {pg}"); r2.font.name="Arial"; r2.font.size=Pt(12)
    _page_break(doc)

    # ── 1. INTRODUÇÃO ─────────────────────────────────────────────────────────
    log("Gerando seções de conteúdo...")
    _heading(doc, "1  INTRODUÇÃO")
    _para(doc, f"Este relatório documenta a execução de um pentest de reconhecimento e varredura sobre o ambiente {domain}, com o objetivo de identificar vulnerabilidades reais por meio de uma abordagem automatizada e adaptativa à tecnologia do alvo.")
    _para(doc, f"Diferentemente de abordagens que assumem uma tecnologia fixa, a ferramenta utilizada neste trabalho identifica automaticamente a stack tecnológica do ambiente-alvo — neste caso, {tech_full} — e adapta os checks de varredura, os endpoints testados e as regras de classificação de vulnerabilidades de acordo com o perfil detectado.")

    # ── 2. IDENTIFICAÇÃO ───────────────────────────────────────────────────────
    _heading(doc, "2  IDENTIFICAÇÃO DO PROJETO")
    id_rows = [
        ["Equipe", team],
        ["Alvo (URL)", f"https://{domain}"],
        ["Tecnologia detectada", tech_full],
        ["Categoria", tech_cat or "N/A"],
        ["CDN / WAF", dns.get("cdn","N/A")],
        ["Metodologia", "OWASP WSTG v4.2"],
        ["Data do relatório", date],
    ]
    if members: id_rows.insert(1, ["Integrantes", " | ".join(members)])
    _add_table(doc, ["Campo","Descrição"], id_rows, [4.5, 11.5])

    # ── 3. METODOLOGIA ──────────────────────────────────────────────────────────
    _heading(doc, "3  METODOLOGIA")
    _para(doc, "A varredura seguiu as seguintes etapas, executadas de forma automatizada:")
    for item in [
        "Reconhecimento passivo: WHOIS, DNS (A, NS, MX, TXT) e verificação de DNSSEC;",
        "Análise de cabeçalhos HTTP: HSTS, CSP, X-Frame-Options, CORS, cookies de segurança;",
        "Fingerprinting de tecnologias: identificação automática da stack via catálogo de assinaturas (HTML, cabeçalhos, cookies);",
        "Enumeração de endpoints: combinação de checks genéricos (painéis admin, arquivos sensíveis universais) com checks específicos da tecnologia detectada;",
        "Varredura de vulnerabilidades: Nikto, verificação de ETag, SSL/TLS, CSRF e rate limiting;",
        "Classificação de vulnerabilidades: análise por IA (Claude/Gemini) com fallback para motor de regras determinístico local.",
    ]:
        _bullet(doc, item)

    # ── 4. RECONHECIMENTO E VARREDURA ───────────────────────────────────────────
    _heading(doc, "4  RECONHECIMENTO E VARREDURA")

    _heading(doc, "4.1  Tecnologia Detectada", level=2)
    matches = tech.get("tech_matches", [])
    if matches:
        match_rows = [[m["name"], m["category"], str(m["score"])] for m in matches[:8]]
        _add_table(doc, ["Tecnologia Candidata","Categoria","Confiança (score)"], match_rows, [6.5, 5.0, 4.5])
        _para(doc, f"A tecnologia principal identificada foi {tech_full}, com maior pontuação de confiança entre as assinaturas avaliadas pelo catálogo de fingerprinting.")
    else:
        _para(doc, "Nenhuma tecnologia conhecida do catálogo foi identificada com confiança suficiente. O ambiente foi tratado como stack genérica, aplicando-se apenas os checks universais de segurança web.")

    _heading(doc, "4.2  DNS e WHOIS", level=2)
    dns_rows = [
        ["Domínio", domain],
        ["IPs (A record)", ", ".join(dns.get("a_records",[])) or "N/A"],
        ["Nameservers", ", ".join(dns.get("ns_records",[])) or "N/A"],
        ["MX", ", ".join(dns.get("mx_records",[])) or "N/A"],
        ["CDN detectado", dns.get("cdn","N/A")],
        ["DNSSEC ativo", "SIM" if dns.get("dnssec_enabled") else "NÃO"],
        ["Registrante", dns.get("registrant","N/A")],
    ]
    _add_table(doc, ["Campo","Valor"], dns_rows, [5.5, 10.5])

    _heading(doc, "4.3  Cabeçalhos HTTP", level=2)
    hdr_rows = []
    for f in hdrs.get("findings",[]):
        sev = f.get("severity","")
        icon = "✔" if sev=="OK" else ("✘" if sev in ("Alto","Médio") else "⚠")
        hdr_rows.append([f["header"], f"{icon} {f['status']}", sev if sev!="OK" else "Conforme"])
    if hdr_rows:
        _add_table(doc, ["Cabeçalho HTTP","Status","Severidade"], hdr_rows, [4.5, 8.0, 3.5])

    _heading(doc, "4.4  Endpoints e Arquivos Sensíveis", level=2)
    ep_rows = []
    for e in ep.get("endpoints",[]):
        st = e.get("status",0)
        if st in (200,301,302,403,500):
            obs = {200:"Acessível (200)",301:"Redirect 301",302:"Redirect 302",403:"Bloqueado (403)",500:"Erro 500"}.get(st,str(st))
            ep_rows.append([e.get("path",""), obs, e.get("label","")])
    if ep_rows:
        _add_table(doc, ["Endpoint","Status","Descrição"], ep_rows[:30], [5.5, 4.0, 6.5])

    sensitive_files = ep.get("sensitive_files_found", [])
    if sensitive_files:
        _para(doc, "Foram identificados os seguintes arquivos sensíveis genéricos publicamente acessíveis:")
        for sf in sensitive_files:
            _bullet(doc, sf["path"])

    _heading(doc, "4.5  Tabela Consolidada de Vulnerabilidades", level=2)
    if vuls:
        vul_rows = [[v.get("id",""), v.get("title",""), v.get("severity",""), v.get("owasp",""), v.get("certainty",""), v.get("tool","")] for v in vuls]
        _add_table(doc, ["ID","Vulnerabilidade","Severidade","OWASP","Certeza","Ferramenta"], vul_rows, [1.2,4.5,1.3,3.0,1.3,4.7])

    # ── 5. ANÁLISE DE IMPACTO ───────────────────────────────────────────────────
    _heading(doc, "5  ANÁLISE DE IMPACTO")
    if vuls:
        imp_rows = [[v.get("id",""), v.get("title",""), v.get("impact","N/A")] for v in vuls]
        _add_table(doc, ["ID","Vulnerabilidade","Impacto Técnico e de Negócio"], imp_rows, [1.2,3.5,11.3])

    # ── 6. SUGESTÕES DE CORREÇÃO ─────────────────────────────────────────────────
    _heading(doc, "6  SUGESTÕES DE CORREÇÃO")
    if vuls:
        rem_rows = [[f"{v.get('id','')} – {v.get('severity','').upper()}", v.get("remediation","N/A")] for v in vuls]
        _add_table(doc, ["ID / Prioridade","Recomendação de Correção"], rem_rows, [3.5,12.5])

    # ── 7. CONSIDERAÇÕES FINAIS ──────────────────────────────────────────────────
    _heading(doc, "7  CONSIDERAÇÕES FINAIS")
    _para(doc, f"O reconhecimento automatizado de {domain} identificou a stack {tech_full} e, a partir dessa identificação, adaptou dinamicamente os endpoints testados e as regras de classificação de vulnerabilidades, resultando em {len(vuls)} achados preliminares.")
    _para(doc, "A abordagem de detecção automática de tecnologia, combinada com checks genéricos universais e checks específicos por stack, permite que esta mesma ferramenta seja reutilizada em alvos com tecnologias completamente diferentes — CMS, frameworks modernos, e-commerce ou APIs — sem necessidade de reconfiguração manual.")
    _page_break(doc)

    # ── REFERÊNCIAS ───────────────────────────────────────────────────────────────
    _section_title(doc, "REFERÊNCIAS")
    refs = [
        "OWASP FOUNDATION. OWASP Top 10:2021. Disponível em: https://owasp.org/Top10/.",
        "OWASP FOUNDATION. OWASP Web Security Testing Guide (WSTG) v4.2. Disponível em: https://owasp.org/www-project-web-security-testing-guide/.",
        "NOAKES-FRUSKIN, R. et al. Nikto Web Scanner. Disponível em: https://github.com/sullo/nikto.",
        "OJ, Oven. Gobuster. Disponível em: https://github.com/OJ/gobuster.",
    ]
    for ref in refs:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing=1.0; p.paragraph_format.space_after=Pt(6)
        p.paragraph_format.left_indent=Cm(1.25); p.paragraph_format.first_line_indent=Cm(-1.25)
        r = p.add_run(ref); r.font.name="Arial"; r.font.size=Pt(12)

    doc.save(str(output_path))
    log(f"Documento salvo: {output_path}")
